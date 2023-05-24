# coding=utf8

'''
实现report自动化

step1:copy the file of report of LY and rename it with new year flag
step2:change parameters in this code and run gen_paras();
step3: revise the year in audit report manually, substitute the above document in the report;
step4: in report, copy the closing to opening manually; write some tables manually;
step5: change parameters and run gen_table_data()

'''
# todo 未自动化部分：待摊费用的发生数；针对每个科目没有实现先清空数据


import warnings

warnings.filterwarnings("ignore")
import os
import re
import datetime
from win32com import client
import pywintypes
import docx  # pip install python-docx
from docx import Document
from docx.oxml.ns import qn  # 字体
from docx.shared import Pt  # 字体大小
from docx.shared import RGBColor  # 颜色
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 对齐
# from docx.enum.text import WD_ALIGN_PARAGRAPH #对齐
from docx.shared import Inches  # 图片大小

import pandas as pd
import numpy as np

# alin with
pd.set_option('display.unicode.ambiguous_as_wide', True)
pd.set_option('display.unicode.east_asian_width', True)
# 显示所有列
pd.set_option('display.max_columns', None)
pd.set_option('display.max_rows', 100)
pd.set_option('max_colwidth', 250)
pd.set_option('display.width', 160)


def gen_paras_report():
    '''
    2022年审计典当协会的附注部分，不含tables
    '''

    print(docx.__version__)
    '''get BS numbers'''
    tb = "典当行业报表及附注202206.xlsm"
    file_path=r"202206\典当行业报表及附注202206.xlsm"
    df_BS = pd.read_excel(rf'D:\work\{file_path}', sheet_name="资产负债表")
    df_IS = pd.read_excel(rf'D:\work\{file_path}', sheet_name="业务活动表", header=4, index_col=0, usecols=[0, 2, 5])
    df_IS = df_IS.drop(np.nan)  # 将索引为nan的删除
    df_IS.index = list(map(lambda x: x.strip(), df_IS.index))  # strip index
    df_IS = df_IS.fillna(0)
    df_CS = pd.read_excel(rf'D:\work\{file_path}', sheet_name="现金流量表")

    '目测一下表清理的怎样，人机结合'
    print(df_IS)

    print(df_BS)
    LU_row = np.where(df_BS == "资产")[0][0]  # left_up
    LU_col = np.where(df_BS == "资产")[1][0]
    RD_col = np.where(df_BS == "负债和净资产合计")[1][0]  # right_down
    RD_row = np.where(df_BS == "负债和净资产合计")[0][0]
    print(LU_row, RD_row, RD_col)
    df_BS = df_BS.iloc[LU_row:RD_row + 1, LU_col:RD_col + (RD_col - LU_col)]
    print("striped BS:\n", df_BS)

    # BS一分为二，并清洗并于索引
    df_BS = df_BS.set_axis(df_BS.iloc[0, :], axis=1)
    df_BS = df_BS.iloc[1:, :]

    df_BS_1 = df_BS.iloc[:, :int(df_BS.shape[1] / 2)]
    df_BS_1 = df_BS_1.dropna(thresh=1)
    df_BS_1 = df_BS_1.set_index('资产', drop=True)
    df_BS_1 = df_BS_1.fillna(0.00)
    df_BS_1 = df_BS_1.applymap(lambda x: round(x, 2))
    df_BS_1.index = list(map(lambda x: x.strip(), df_BS_1.index))
    df_BS_2 = df_BS.iloc[:, int(df_BS.shape[1] / 2):df_BS.shape[1]]
    df_BS_2 = df_BS_2.dropna(thresh=1)
    df_BS_2 = df_BS_2.set_index('负债和净资产', drop=True)
    df_BS_2 = df_BS_2.fillna(0.00)
    df_BS_2 = df_BS_2.applymap(lambda x: round(x, 2))
    df_BS_2.index = list(map(lambda x: x.strip(), df_BS_2.index))

    print("assets:\n", df_BS_1)
    print("liabilities:\n", df_BS_2)

    print(df_BS_1.loc["货币资金", "年初数"])

    '''docx edit audit_reports'''

    file = docx.Document()
    file.styles['Normal'].font.name = u'宋体'
    file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    head = file.add_heading('审计报告', level=0)
    head.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # head.add_run()
    file.add_paragraph()
    file.add_paragraph()
    file.add_paragraph()
    #
    #
    # year = "2022"
    date = "2022年6月30日";period = "2022年1-6月"

    assets = 1111.11;
    current_assets = 222;
    L_invest = 333;
    fixed = 444;
    depriation = 555;
    net_fixed = 666;
    intangible = 777;
    proxy_a = 888
    paras = ["六、资产状况",
             f"1、截止{date}，SCSDDHYXH资产总额为{df_BS_1.loc['资产总计', '年末数']:,.2f}元，其中："
             f"流动资产{df_BS_1.loc['流动资产合计', '年末数']:,.2f}元，长期投资{df_BS_1.loc['长期投资合计', '年末数']:,.2f}元；"
             f"固定资产原值{df_BS_1.loc['固定资产原价', '年末数']:,.2f}元，累计折旧{df_BS_1.loc['减：累计折旧', '年末数']:,.2f}元，"
             f"固定资产净值{df_BS_1.loc['固定资产净值', '年末数']:,.2f}元，无形资产{df_BS_1.loc['无形资产', '年末数']:,.2f}元，"
             f"受托代理资产{df_BS_1.loc['受托代理资产', '年末数']:,.2f}元。",
             f"2、截止{date}，SCSDDHYXH负债总额为{df_BS_2.loc['负债合计', '年末数']:,.2f}元，其中："
             f"流动负债{df_BS_2.loc['流动负债合计', '年末数']:,.2f}元,长期负债{df_BS_2.loc['长期负债合计', '年末数']:,.2f}元，"
             f"受托代理负债{df_BS_2.loc['受托代理负债', '年末数']:,.2f}元。",
             f"3、截止{date}，SCSDDHYXH净资产总额{df_BS_2.loc['净资产合计', '年末数']:,.2f}元，其中："
             f"限定性净资产{df_BS_2.loc['限定性净资产', '年末数']:,.2f}元，非限定性净资产{df_BS_2.loc['非限定性净资产', '年末数']:,.2f}元。",
             "七、收支情况",
             f"1、SCSDDHYXH{period}收入{df_IS.loc['收入合计', '本年累计数']:,.2f}元，"
             f"其中：捐赠收入{df_IS.loc['其中：捐赠收入', '本年累计数']:,.2f}元，会费收入{df_IS.loc['会费收入', '本年累计数']:,.2f}元，"
             f"提供服务收入{df_IS.loc['提供服务收入', '本年累计数']:,.2f}元，商品销售收入{df_IS.loc['商品销售收入', '本年累计数']:,.2f}元，"
             f"政府补贴收入{df_IS.loc['政府补助收入', '本年累计数']:,.2f}元，投资收益{df_IS.loc['投资收益', '本年累计数']:,.2f}元，"
             f"其他收入{df_IS.loc['其他收入', '本年累计数']:,.2f}元。",
             f"2、SCSDDHYXH{period}费用{df_IS.loc['费用合计', '本年累计数']:,.2f}元，"
             f"其中：业务活动成本{df_IS.loc['（一）业务活动成本', '本年累计数']:,.2f}元，管理费用{df_IS.loc['（二）管理费用', '本年累计数']:,.2f}元，"
             f"筹资费用{df_IS.loc['（三）筹资费用', '本年累计数']:,.2f}元，其他费用{df_IS.loc['（四）其他费用', '本年累计数']:,.2f}元。",
             "====================",
             "七、净资产变动额：",
             f"{period}净资产变动额为{df_IS.loc['四、净资产变动额（若为净资产减少额，以“-”号填列）', '本年累计数']:,.2f}元。收入合计{df_IS.loc['收入合计', '本年累计数']:,.2f}元,"
             f" 费用合计{df_IS.loc['费用合计', '本年累计数']:,.2f}元，净资产变动额为收入合计减去费用合计，"
             f"净资产增加{df_IS.loc['四、净资产变动额（若为净资产减少额，以“-”号填列）', '本年累计数']:,.2f}元。",
             f"资产负债表净资产年初数{df_BS_2.loc['净资产合计', '年初数']:,.2f}元，年末数{df_BS_2.loc['净资产合计', '年末数']:,.2f}元，净资产增加{df_IS.loc['四、净资产变动额（若为净资产减少额，以“-”号填列）', '本年累计数']:,.2f}元。",
             "以下为附送8内容",
             "二、资产情况",
             f"截至{date}，SCSDDHYXH资产总额为{df_BS_1.loc['资产总计', '年末数'] / 10000:,.2f}万元，"
             f"负债总额为{df_BS_2.loc['负债合计', '年末数'] / 10000:,.2f}万元，净资产总额为{df_BS_2.loc['净资产合计', '年末数'] / 10000:,.2f}万元，"
             f"其中限定性净资产为0.00万元，非限定性净资产为{df_BS_2.loc['净资产合计', '年末数'] / 10000:,.2f}万元。",
             "（一）结构及变动情况",
             f"截至{date}，SCSDDHYXH资产总额为{df_BS_1.loc['资产总计', '年末数'] / 10000:,.2f}万元，"
             f"较上年度增加{df_BS_1.loc['资产总计', '年末数'] / 10000 - df_BS_1.loc['资产总计', '年初数'] / 10000:,.2f}万元，"
             f"其中流动资产{df_BS_1.loc['流动资产合计', '年末数'] / 10000:,.2f}万元，较上年度增加{df_BS_1.loc['流动资产合计', '年末数'] / 10000 - df_BS_1.loc['流动资产合计', '年初数'] / 10000:,.2f}万元；"
             f"长期投资0.00万元，较上年度增加0.00万元；固定资产{df_BS_1.loc['固定资产净值', '年末数'] / 10000:,.2f}万元，较上年度减少{df_BS_1.loc['固定资产净值', '年初数'] / 10000 - df_BS_1.loc['固定资产净值', '年末数'] / 10000:,.2f}万元；"
             f"无形资产0.00万元，较上年度增0.00万元；受托代理资产0.00万元，较上年度增加0.00万元。",
             f"截至{date}， SCSDDHYXH负债总额为{df_BS_2.loc['负债合计', '年末数'] / 10000:,.2f}万元，较上年度增加{df_BS_2.loc['负债合计', '年末数'] / 10000 - df_BS_2.loc['负债合计', '年初数'] / 10000:,.2f}万元，"
             f"其中流动负债{df_BS_2.loc['流动负债合计', '年末数'] / 10000:,.2f}万元，较上年度增加{df_BS_2.loc['流动负债合计', '年末数'] / 10000 - df_BS_2.loc['流动负债合计', '年初数'] / 10000:,.2f}万元；"
             f"长期负债0.00万元，较上年度增加0.00万元；受托代理负债0.00万元，较上年度增加0.00万元。",
             f"截至{date}，SCSDDHYXH净资产总额为{df_BS_2.loc['净资产合计', '年末数'] / 10000:,.2f}万元，"
             f"较上年度增加{df_BS_2.loc['净资产合计', '年末数'] / 10000 - df_BS_2.loc['净资产合计', '年初数'] / 10000:,.2f}万元。",
             "（二）变动原因分析",
             "分别详细说明“（一）结构及变动情况”中有增减变化的资产和负债原因。",
             f"流动资产减少32.83万元主要系货币资金支付；",
             f"流动负债减少1.50万元主要系预收结转。"]

    for i, v in enumerate(paras[:]):
        p = file.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
        p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
        p.paragraph_format.line_spacing = Pt(26)  # 行距
        p.paragraph_format.space_before = Pt(0)  # 段前间距
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(v)
        font = run.font
        if re.match("[六|七|八|九|二|（一）结|（二）变]\\b、", v):
            print(i)
            # p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
            # p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
            # p.paragraph_format.line_spacing = Pt(26)  # 行距
            # p.paragraph_format.space_before = Pt(0)  # 段前间距
            # p.paragraph_format.space_after = Pt(0)
            # run = p.add_run(v)
            # font = run.font
            font.bold = True
            font.size = Pt(12)  # 小四
        else:
            # p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
            # p.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
            # p.paragraph_format.line_spacing = Pt(26)  # 行距
            # p.paragraph_format.space_before = Pt(0)  # 段前间距
            # p.paragraph_format.space_after = Pt(0)
            # run = p.add_run(v)
            # font = run.font
            font.size = Pt(10.5)
        font.color.rgb = RGBColor(0, 0, 0)

    # add pictures
    # file.add_picture(f"D:\python\sealed\concepts\\money_flow.png", width=None, height=None)
    # file.add_paragraph("second")
    # file.add_picture(f"D:\python\sealed\concepts\\money_flow.png", width=Inches(10), height=Inches(5))
    # file.add_paragraph("third")
    # file.add_picture(f"D:\python\sealed\concepts\\money_flow.png", width=Inches(11), height=Inches(6))

    # todo need to close at first
    file.save(f"D:\work\\202206\\audit_report2206.docx")
    os.startfile(f"D:\work\\202206\\audit_report2206.docx")

def cls_opn():
    '''
    实现附注中上年数的copy
    :return:
    '''
    pass


def gen_table_data(file,acc,df_acc):
    """
    部分科目，实现附注中table数据本年数的填写
    src: 余额表；tgt: word中的table
    :return:
    """
    # find src

    acc_code = df_acc.loc[acc, "科目编码"]
    print(acc_code)
    df_detail = df_acc[df_acc["科目编码"].str.contains(re.compile(f"{acc_code}\d*"), regex=True)]
    df_detail.index = list(map(lambda x:x.strip(),df_detail.index))
    print(df_detail)
    # todo check whether the total is correct,防止有虚增发生额情形


    # step2: find the target tale in word

    tables = file.tables
    print("total tables:", len(tables))

    def find_table(str):
        i = 0;
        j = 0;
        k = 0
        for table in file.tables:
            j = 0
            for row in table.rows:
                k = 0
                for cell in row.cells:
                    if cell.text == str:
                        print("got it")
                        return (i, j, k)
                    k += 1
                j += 1
            i += 1

    print(flags.get(acc))
    position = find_table(flags.get(acc))
    print(position)
    table_tgt = tables[position[0]]
    print("table_tgt:\n", table_tgt)

    # 3 substitute from src to tgt
    # 3.1 check whether we need to add a row
    # print(len(table_tgt.rows)) # total 15 rows in word including 2 headers and 1 total
    # print(len(df_detail)) # total 12 rows in acc including 1 general ledger.
    # 根据word中的明细科目找寻df中的数据然后填列
    # 根据科目的性质，判断需要填列table中的那些列，例如损益类需要填第0，1，3列
    if int(acc_code[:1])>3:  # 损益类科目
        cols_fillin = [0, 1, 3]
        if re.match("5301$",acc_code):
            cols_fillin=[0,2,3]

        col_copy = "本期发生借方"
    elif int(acc_code[:1])==1: # 资产类科目
        # 资产负债类科目应为含有“末”的列号
        lst=[i for i in range(len(table_tgt.columns)) if re.search("末",table_tgt.rows[0].cells[i].text)]
        cols_fillin = [0]+lst
        col_copy = "期末借方"
    else:  # 负债权益类科目
        # 资产负债类科目应为含有“末”的列号
        lst = [i for i in range(len(table_tgt.columns)) if re.search("末", table_tgt.rows[0].cells[i].text)]
        cols_fillin = [0] + lst
        col_copy = "期末贷方"



        pass

    print(cols_fillin)
    for row in table_tgt.rows[:]:
        if row.cells[0].text=="项目":continue
        acc_tier2 = row.cells[0].text

        for j in cols_fillin[1:]:
            if not acc_tier2 in df_detail.index: # 上年有本年无的明细科目
                if acc_tier2=="合计":
                    temp = df_detail.loc[acc,col_copy]
                else:
                    temp =0.00
            else:
                temp = df_detail.loc[acc_tier2, col_copy]

            if temp!=0: row.cells[j].text=f"{temp:,.2f}"
            else: row.cells[j].text=""
            row.cells[j].paragraphs[
                0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 需要先录入text,再修改格式，否则paragraph找不到



            print(row.cells[j].text)

    # todo 检查是否录入DF完毕，是否有DF明细科目大于WORD中table明细的



    # print(table_tgt.cell(7, 0).text)
    # table_tgt.cell(7, 0).text = "社保费"
    # a = 3000
    #





    pass


if __name__ == '__main__':

    # gen_paras_report()

    cls_opn()

    file_path2=r"202206\202206典当行业财务审计报告.docx"
    file = docx.Document(rf"d:\work\{file_path2}")
    # 当在word中定位一个会计科目的明细表table时，需要使用到这个科目的明细作为table的一个标志
    # 注意word中明细与科目余额表一直


    df_acc = pd.read_excel(r'd:\work\202206\余额表202206.xls', usecols=(2, 3, 5, 6, 7, 8, 9, 10),
                       index_col="科目名称")
    # print(df_acc)
    df_acc = df_acc[list(map(lambda x:isinstance(x,str),df_acc.index))]
    df_acc.index = list(map(lambda x: x.strip(), df_acc.index))

    # AJE 对科目余额表的影响
    df_acc.loc["其他收入"] = ["420102", 0, 0, 4600.16, 0, 0, 0]
    df_acc.loc["利息收入", "本期发生借方"] = df_acc.loc["利息收入", "本期发生借方"] + 4600.16
    df_acc.loc["财务费用", "本期发生借方"] = df_acc.loc["财务费用", "本期发生借方"] + 4600.16
    df_acc.rename(index={"财务费用":"其他费用","手续费":"银行手续费"},inplace=True)
    df_acc.loc["租房费用"]=["520108"]+[0 for i in range(len(df_acc.columns)-1)]
    df_acc.loc["租房费用", "本期发生借方"]=112446
    df_acc.loc["管理费用", "本期发生借方"] = df_acc.loc["管理费用", "本期发生借方"]+112446
    # print(df_acc.loc["待摊费用", "期末借方"])

    df_acc.loc["待摊费用", "期末借方"]=df_acc.loc["待摊费用", "期末借方"]-112446
    df_acc.loc["房租", "期末借方"]=df_acc.loc["房租", "期末借方"]-112446
    # print(df_acc.loc["待摊费用", "期末借方"])
    print(df_acc)

    # 根据附注披露要求，修正科目余额表，方便运行代码 5101 会员服务成本
    df_acc=df_acc[~df_acc["科目编码"].str.contains(re.compile(f"5101\d+"), regex=True)]
    df_acc.loc["会员服务成本"]=["510199"]+list(df_acc.loc["业务活动成本"][1:])
    df_acc.loc["收入"]=["4201"]+[0 for i in range(len(df_acc.columns)-1)]
    df_acc.loc["收入","本期发生借方"]=df_acc.loc["会费收入","本期发生借方"]+df_acc.loc["其他收入","本期发生借方"]
    df_acc.loc["会费收入", "科目编码"]="420101";df_acc.loc["其他收入", "科目编码"]="420102"
    df_acc.loc["货币资金"] = ["1001"] + [0 for i in range(len(df_acc.columns) - 1)]
    df_acc.loc["货币资金", "期末借方"] = df_acc.loc["现金", "期末借方"] + df_acc.loc["银行存款", "期末借方"]
    df_acc.loc["现金", "科目编码"] = "100101";df_acc.loc["银行存款", "科目编码"] = "100102"
    # print(df_acc)
    df_acc.loc["工资费用"] = ["5301"] + [0 for i in range(len(df_acc.columns) - 1)]
    df_acc.loc["一、工资、奖金、津贴和补贴"] = ["530101"] + list(df_acc.loc["工资"][1:])
    df_acc.loc["二、职工福利费"] = ["530102"] + list(df_acc.loc["福利费"][1:])
    df_acc.loc["三、社会保险费"] = ["530103"] + list(df_acc.loc["社保费"][1:])
    # 修改其中一个重复的索引名，怎么这么麻烦
    idx = np.where(df_acc.index=="住房公积金")[0][0]
    lst=list(df_acc.index)
    lst[idx]="qt-住房公积金"
    df_acc.index=lst
    # print(df_acc.index)
    df_acc.loc["四、住房公积金"] = ["530104"] + list(df_acc.loc["住房公积金"][1:])
    df_acc.loc["工资费用", "本期发生借方"] =df_acc[df_acc["科目编码"].str.contains(re.compile(f"5301\d+"), regex=True)]["本期发生借方"].sum()
    print(df_acc)
    df_acc.loc["净资产"] = ["3101"] + [0 for i in range(len(df_acc.columns) - 1)]
    df_acc.loc["净资产", "期末贷方"] = df_acc.loc["非限定性净资产", "期末贷方"]
    df_acc.loc["非限定性净资产", "科目编码"] = "310101"
    print(df_acc)




    flags = {"管理费用":"租房费用","其他费用":"银行手续费","业务活动成本":"会员服务成本",
             "收入":"会费收入","货币资金":"银行存款","待摊费用":"房租","应交税金":"应交代扣个人所得税",
             "工资费用":"一、工资、奖金、津贴和补贴","净资产":"非限定性净资产"}
    # 当明细名称在科目余额表同附注披露不一致时，可以改动附注，这样后续年度省事
    # todo 其他应收款由于本期无变化，未修改实用程序

    for key,value in flags.items():
        print()
        gen_table_data(file,key,df_acc)



    file.save(rf"d:\work\{file_path2}")
